import os
import re
from docx import Document
from faker import Faker

from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig


fake = Faker(["en_IN", "en_US"])


ENTITY_MAPPING_CACHE = {}


def get_synthetic_replacement(entity_type: str, original_val: str) -> str:
    """Retrieves or generates a consistent fake replacement for PII entities."""
    key = (entity_type, original_val.strip())
    if key in ENTITY_MAPPING_CACHE:
        return ENTITY_MAPPING_CACHE[key]

    if entity_type == "PERSON":
        replacement = fake.name()
    elif entity_type == "EMAIL_ADDRESS":
        replacement = fake.email()
    elif entity_type == "PHONE_NUMBER":
        replacement = fake.phone_number()
    elif entity_type in ["LOCATION", "ADDRESS"]:
        replacement = fake.address().replace("\n", ", ")
    elif entity_type == "ORGANIZATION":
        replacement = fake.company()
    elif entity_type == "US_SSN":
        replacement = fake.ssn()
    elif entity_type == "IN_PAN":
        replacement = f"[REDACTED_PAN_{fake.bothify(text='?????####?')}]"
    elif entity_type == "IN_AADHAAR":
        replacement = f"[REDACTED_AADHAAR_{fake.numerify(text='#### #### ####')}]"
    elif entity_type == "CREDIT_CARD":
        replacement = fake.credit_card_number()
    elif entity_type == "DATE_TIME":
        replacement = fake.date_of_birth().strftime("%Y-%m-%d")
    elif entity_type == "IP_ADDRESS":
        replacement = fake.ipv4()
    else:
        replacement = f"[REDACTED_{entity_type}]"

    ENTITY_MAPPING_CACHE[key] = replacement
    return replacement


def setup_analyzer_engine() -> AnalyzerEngine:
    """Configures Presidio Analyzer with custom Indian PII recognizers."""
    analyzer = AnalyzerEngine()

    
    pan_pattern = Pattern(
        name="pan_pattern",
        regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b",
        score=0.95,
    )
    pan_recognizer = PatternRecognizer(
        supported_entity="IN_PAN", patterns=[pan_pattern]
    )

    
    aadhaar_pattern = Pattern(
        name="aadhaar_pattern",
        regex=r"\b[2-9]{1}[0-9]{3}\s[0-9]{4}\s[0-9]{4}\b",
        score=0.95,
    )
    aadhaar_recognizer = PatternRecognizer(
        supported_entity="IN_AADHAAR", patterns=[aadhaar_pattern]
    )

    analyzer.registry.add_recognizer(pan_recognizer)
    analyzer.registry.add_recognizer(aadhaar_recognizer)
    return analyzer


def redact_text_block(
    text: str, analyzer: AnalyzerEngine, anonymizer: AnonymizerEngine
) -> str:
    """Detects PII in string blocks and replaces them with synthetic data."""
    if not text or not text.strip():
        return text

    target_entities = [
        "PERSON",
        "EMAIL_ADDRESS",
        "PHONE_NUMBER",
        "LOCATION",
        "ORGANIZATION",
        "US_SSN",
        "IN_PAN",
        "IN_AADHAAR",
        "CREDIT_CARD",
        "DATE_TIME",
        "IP_ADDRESS",
    ]

    results = analyzer.analyze(
        text=text, entities=target_entities, language="en", score_threshold=0.4
    )

    if not results:
        return text

    operators = {}
    for res in results:
        entity_val = text[res.start : res.end]
        replacement = get_synthetic_replacement(res.entity_type, entity_val)
        operators[res.entity_type] = OperatorConfig(
            "replace", {"new_value": replacement}
        )

    anonymized_res = anonymizer.anonymize(
        text=text, analyzer_results=results, operators=operators
    )

    return anonymized_res.text


def process_docx(input_file: str, output_file: str):
    """Processes document paragraphs and tables to produce a redacted Word document."""
    doc = Document(input_file)
    analyzer = setup_analyzer_engine()
    anonymizer = AnonymizerEngine()

    
    for p in doc.paragraphs:
        if p.text.strip():
            redacted = redact_text_block(p.text, analyzer, anonymizer)
            if redacted != p.text:
                p.text = redacted

   
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        redacted = redact_text_block(
                            p.text, analyzer, anonymizer
                        )
                        if redacted != p.text:
                            p.text = redacted

    doc.save(output_file)
    print(f"Redaction completed successfully! File saved to {output_file}")


if __name__ == "__main__":
    input_doc = "Red Herring Prospectus (1).docx"
    output_doc = "Redacted_Red_Herring_Prospectus.docx"
    if os.path.exists(input_doc):
        process_docx(input_doc, output_doc)